"""比对 g2t 管道输出与文献参考数据。

用法:
  python scripts/compare_results.py
"""
import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document

BASE = Path("/media/deyuan/217ce44c-b5de-45ed-8720-deebdff85ece/CLAUDE/ceshi")


# ---------------------------------------------------------------------------
# Part 1: Opheliidae 新旧输出比对
# ---------------------------------------------------------------------------
def compare_ophelidae():
    old = BASE / "ceshi4-ophelidae/output2"
    new = BASE / "ceshi4-ophelidae/output3"
    out = BASE / "ophelidae_diff.csv"

    pairs = [
        ("gb_metadata/final.csv", "final"),
        ("labeled_genes/assigned_genes_types_all.csv", "assigned"),
        ("organized_genes/organized_species_voucher.csv", "organized"),
    ]

    rows = []
    for rel, label in pairs:
        fp_old = old / rel
        fp_new = new / rel
        if not fp_old.exists() or not fp_new.exists():
            rows.append(dict(file=label, status="文件不存在", detail=f"old={fp_old.exists()}, new={fp_new.exists()}"))
            continue

        df_old = pd.read_csv(fp_old, low_memory=False)
        df_new = pd.read_csv(fp_new, low_memory=False)

        if df_old.shape != df_new.shape:
            rows.append(dict(file=label, status="行列数不同", detail=f"old={df_old.shape}, new={df_new.shape}"))

        # 按关键列排序后逐行比较
        key_col = "LocusID" if "LocusID" in df_old.columns else "species_voucher_new"
        if key_col in df_old.columns and key_col in df_new.columns:
            df_old = df_old.sort_values(key_col).reset_index(drop=True)
            df_new = df_new.sort_values(key_col).reset_index(drop=True)

        # 比较共有列
        common_cols = list(set(df_old.columns) & set(df_new.columns))
        for col in common_cols:
            old_vals = df_old[col].fillna("").astype(str).tolist()
            new_vals = df_new[col].fillna("").astype(str).tolist()
            for i, (ov, nv) in enumerate(zip(old_vals, new_vals)):
                if ov != nv:
                    key_val = df_old.at[i, key_col] if key_col in df_old.columns else i
                    rows.append(dict(file=label, column=col, key=key_val,
                                     old_value=ov[:200], new_value=nv[:200]))

    if rows:
        pd.DataFrame(rows).to_csv(out, index=False)
        print(f"Opheliidae 差异: {len(rows)} 条 -> {out}")
    else:
        print("Opheliidae: 新旧输出完全一致，无差异")
        # 写空 CSV 表明一致
        pd.DataFrame(columns=["file", "column", "key", "old_value", "new_value"]).to_csv(out, index=False)


# ---------------------------------------------------------------------------
# Part 2: Hesionidae g2t 输出 vs 文献 docx
# ---------------------------------------------------------------------------
GENE_MAP = {"COI": "coi", "16S": "16s", "18S": "18s", "28S": "28s"}


def parse_docx(path):
    """解析 zlag028_supplementary_data.docx，返回 DataFrame[species, gene, accession]"""
    doc = Document(path)
    table = doc.tables[0]

    data = []
    for row in table.rows[1:]:  # 跳过表头
        cells = [c.text.strip() for c in row.cells]
        species_raw = cells[0]
        refs = cells[5] if len(cells) > 5 else ""
        for i, gene_label in enumerate(["COI", "16S", "18S", "28S"]):
            accession = cells[i + 1].strip() if len(cells) > i + 1 else ""
            if accession == "–" or accession == "-" or accession == "":
                accession = ""
            data.append(dict(
                species_raw=species_raw,
                gene=gene_label,
                literature_accession=accession,
                references=refs,
            ))
    return pd.DataFrame(data)


def normalize_species(name):
    """将物种名标准化为 lowercase，去除 author/year 等"""
    name = name.strip()
    # 去除 "sp. nov.", "sp.nov.", "sensu ...", "cf." 等修饰词
    name = re.sub(r"\s*(sp\.?\s*nov\.?|sensu\s+.*|cf\.\s*)", " ", name, flags=re.I)
    # 去除 author year (如 "(Lamarck, 1818)" 或 "Jimi, 2018")
    name = re.sub(r"\([^)]*\d{4}[^)]*\)", "", name)
    name = re.sub(r",?\s*[A-Z][a-z]+,?\s*\d{4}.*$", "", name)
    # 去除 "sp. A", "sp. D" 等保留
    name = re.sub(r"\s+", " ", name).strip()
    return name.lower()


def match_species(lit_species, g2t_organisms):
    """在 g2t organism 列表中找到最佳匹配。返回匹配的 organism 值或 None"""
    lit_norm = normalize_species(lit_species)
    # 精确匹配
    for org in g2t_organisms:
        if normalize_species(org) == lit_norm:
            return org
    # 模糊: 去除 sp. 后的前缀匹配
    lit_prefix = re.sub(r"\s+sp\.?\s*[A-Z]?$", "", lit_norm).strip()
    for org in g2t_organisms:
        org_prefix = re.sub(r"\s+sp\.?\s*[A-Z]?$", "", normalize_species(org)).strip()
        if lit_prefix and org_prefix and (lit_prefix == org_prefix or lit_prefix in org_prefix or org_prefix in lit_prefix):
            return org
    # 尝试属名+种名前几个字符
    lit_parts = lit_norm.split()
    if len(lit_parts) >= 2:
        lit_genus_sp = f"{lit_parts[0]} {lit_parts[1][:4]}"
        for org in g2t_organisms:
            org_parts = normalize_species(org).split()
            if len(org_parts) >= 2 and f"{org_parts[0]} {org_parts[1][:4]}" == lit_genus_sp:
                return org
    return None


def compare_hesionidae():
    docx_path = BASE / "ceshi3-Hesionidae/zlag028_supplementary_data.docx"
    g2t_assigned = BASE / "ceshi3-Hesionidae/output/labeled_genes/assigned_genes_types_all.csv"
    out = BASE / "hesionidae_vs_literature.csv"

    # 1. 解析 docx 为每物种一行的宽表
    doc = Document(docx_path)
    table = doc.tables[0]
    genes = ["COI", "16S", "18S", "28S"]

    lit_rows = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        species = cells[0]
        refs = cells[5] if len(cells) > 5 else ""
        gene_accs = {}
        for i, g in enumerate(genes):
            acc = cells[i + 1].strip() if len(cells) > i + 1 else ""
            if acc in ("–", "-", ""):
                acc = ""
            gene_accs[f"lit_{GENE_MAP[g]}"] = acc
        lit_rows.append(dict(species=species, **gene_accs, references=refs))

    lit_df = pd.DataFrame(lit_rows)
    print(f"文献数据: {len(lit_df)} 个物种")

    # 2. 读取 g2t assigned 表，按 organism + gene_type 聚合 LocusID
    assigned_df = pd.read_csv(g2t_assigned, low_memory=False)
    agg = (assigned_df.groupby(["organism", "gene_type"])["LocusID"]
           .apply(lambda x: "; ".join(str(v) for v in x.dropna()))
           .reset_index())
    locus_lookup = {}
    for _, r in agg.iterrows():
        locus_lookup[(r["organism"], r["gene_type"])] = r["LocusID"]

    g2t_organisms = assigned_df["organism"].dropna().unique().tolist()

    # 3. 物种匹配 + 逐行比对
    results = []
    for _, row in lit_df.iterrows():
        sp = row["species"]
        matched_org = match_species(sp, g2t_organisms)

        record = {"species": sp, "g2t_organism": matched_org or "未匹配"}
        all_match = True
        has_any_lit = False

        for g in genes:
            gcol = GENE_MAP[g]
            lit_acc = row.get(f"lit_{gcol}", "")
            g2t_acc = locus_lookup.get((matched_org, gcol), "") if matched_org else ""

            record[f"lit_{gcol}"] = lit_acc
            record[f"g2t_{gcol}"] = g2t_acc

            if lit_acc:
                has_any_lit = True
                if g2t_acc and lit_acc in g2t_acc:
                    pass  # 一致
                elif g2t_acc:
                    all_match = False
                else:
                    all_match = False
            elif g2t_acc:
                pass  # 文献无，g2t有，不算不一致

        if not has_any_lit:
            record["match"] = "文献无数据"
        elif matched_org is None:
            record["match"] = "g2t未匹配到该物种"
        elif all_match:
            record["match"] = "一致"
        else:
            record["match"] = "不一致"

        record["references"] = row.get("references", "")
        results.append(record)

    result_df = pd.DataFrame(results)
    result_df.to_csv(out, index=False)

    print(f"\n输出: {len(result_df)} 行 -> {out}")
    print(result_df["match"].value_counts().to_string())

    # 未匹配物种
    unmatched = result_df[result_df["match"] == "g2t未匹配到该物种"]["species"].tolist()
    if unmatched:
        print(f"\ng2t 未匹配到的物种 ({len(unmatched)}):")
        for sp in unmatched:
            print(f"  - {sp}")


if __name__ == "__main__":
    print("=" * 60)
    print("Part 1: Opheliidae 新旧输出比对")
    print("=" * 60)
    compare_ophelidae()

    print("\n" + "=" * 60)
    print("Part 2: Hesionidae g2t vs 文献比对")
    print("=" * 60)
    compare_hesionidae()
